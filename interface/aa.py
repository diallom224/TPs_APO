from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Créer un nouveau document
doc = Document()

# Définir les marges
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Titre principal
title = doc.add_paragraph()
title_run = title.add_run('CHECKLIST PROFIL LINKEDIN')
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Fonction pour créer un tableau de questionnaire
def add_question_section(doc, title, questions):
    # Ajouter le titre de section
    section_title = doc.add_paragraph()
    section_run = section_title.add_run(title)
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_run.font.color.rgb = RGBColor(0, 70, 120)
    
    # Créer le tableau
    table = doc.add_table(rows=len(questions)+1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # En-têtes
    headers = ['QUESTIONS', 'NON', 'OUI', 'Commentaire']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Remplir les questions
    for idx, q in enumerate(questions, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = q['question']
        row_cells[1].text = q.get('non', '')
        row_cells[2].text = q.get('oui', '')
        row_cells[3].text = q.get('commentaire', '')
    
    doc.add_paragraph()

# Section 1 : Photo
questions_photo = [
    {'question': "J'ai une photo", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai fait une photo qui me représente actuellement", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai une tenue adéquate, appropriée et professionnelle", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai un fond neutre", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai une bonne posture", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "Je ne fais pas de grimace", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "Je suis seul sur la photo", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "Je n'ai pas abusé de l'IA pour les modifications (exemple avec pfpmaker)", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "Je souris sur la photo", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'PHOTO DE PROFIL', questions_photo)

# Section 2 : Bandeau
questions_bandeau = [
    {'question': "J'ai un bandeau", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai un bandeau qui me représente actuellement (passion, travail, étude)", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai le droit à l'image du bandeau", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'BANDEAU', questions_bandeau)

# Section 3 : Titre et Résumé
questions_titre = [
    {'question': "J'ai un titre", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "Je n'ai pas fait de faute d'orthographe", 'non': '', 'oui': 'X', 'commentaire': 'À priori pas de faute....'},
    {'question': "J'ai des mots clés me définissant", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai mis ce que je recherche", 'non': '', 'oui': 'X', 'commentaire': 'Faut revoir la partie «résumé/infos»'},
    {'question': "Je n'ai pas mis de propos inappropriés", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai rajouté le résumé", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai rajouté, sur mon profil, une section «info»", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai été claire sur ce que je veux", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai été claire sur ce que je ne veux pas", 'non': 'X', 'oui': '', 'commentaire': ''},
    {'question': "J'ai noté mes coordonnées personnelles", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'TITRE ET RÉSUMÉ', questions_titre)

# Section 4 : Disponibilité
questions_disponibilite = [
    {'question': "J'ai noté ma disponibilité géographique", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai noté ma disponibilité temporelle", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'DISPONIBILITÉ', questions_disponibilite)

# Section 5 : Compétences
questions_competences = [
    {'question': "J'ai mis mon activité professionnelle (si il y en a)", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mis mon activité étudiante (si il y en a)", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mis d'autres types d'activité (associative, bénévolat)", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'COMPÉTENCES', questions_competences)

# Section 6 : Activité LinkedIn
questions_activite = [
    {'question': "Est-ce qu'il y a de l'activité (des posts sur LinkedIn)", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'ACTIVITÉ SUR LINKEDIN', questions_activite)

# Section 7 : Localisation
questions_localisation = [
    {'question': "J'ai bien indiqué la grande ville la plus proche de ma localisation", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai la possibilité de me déplacer dans une autre ville", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'LOCALISATION', questions_localisation)

# Section 8 : Recommandations
questions_recommandations = [
    {'question': "J'ai une recommandation de la part de l'entreprise ou de mon chef de projet", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai une recommandation de la part de mon maître de stage", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'RECOMMANDATIONS', questions_recommandations)

# Section 9 : Astuces
questions_astuces = [
    {'question': "J'ai personnalisé mon URL de profil avec mon nom", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mentionné des projets spécifiques, professionnels ou non", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai enrichi mon profil par des contenus médias", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai dupliqué mon profil dans plusieurs langues si je cherche en dehors de la France", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai activé la fonction #OpenToWork", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai ajouté des informations de contact", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai repartagé des posts avec des avis", 'non': 'X', 'oui': '', 'commentaire': ''},
]
add_question_section(doc, 'ASTUCES UTILES', questions_astuces)

# Section 10 : Catégories d'emploi
questions_categories = [
    {'question': "J'interviens dans le domaine de l'informatique", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai précisé une catégorie pour laquelle j'ai des compétences", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai précisé une catégorie qui me plaît", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'CATÉGORIES D\'EMPLOI', questions_categories)

# Section 11 : Expériences
questions_experiences = [
    {'question': "J'ai mis toutes les expériences professionnelles acquises", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mis toutes les expériences associatives acquises", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mentionné mon rôle et ma mission", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mentionné les compétences acquises", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai notifié l'établissement dans lequel j'ai effectué mon travail", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à marquer tous les outils pour lesquels j'ai acquis une expérience", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à mettre un lien vers mon portfolio si j'en ai un", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à mettre un ou plusieurs liens vers GitHub ou autres", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'EXPÉRIENCES', questions_experiences)

# Section 12 : Formations
questions_formations = [
    {'question': "J'ai indiqué mes formations", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mentionné les notions ou langages étudiés", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai mentionné les compétences acquises", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai notifié l'établissement dans lequel j'ai effectué mes études", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à marquer mes certifications si j'en ai", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à marquer tous les outils pour lesquels j'ai acquis une expérience", 'non': '', 'oui': 'X', 'commentaire': ''},
    {'question': "J'ai pensé à lister les UE correspondantes", 'non': '', 'oui': 'X', 'commentaire': ''},
]
add_question_section(doc, 'FORMATIONS', questions_formations)

# Sauvegarder le document
doc.save('Checklist_LinkedIn.docx')
print("Document créé avec succès : Checklist_LinkedIn.docx")